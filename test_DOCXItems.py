import os
import unittest
from zipfile import ZipFile

from bs4 import BeautifulSoup

from docx.document import DOCXDocument
from docx.document import DOCX_CONTENTS_FILE_NAME, DOCX_RELS_FILE_NAME
from docx.items import DOCXBr, DOCXRun, DOCXText
from docx.items import DOCXDrawing, DOCXHyperlink, DOCXParagraph

DEFAULT_PARSER = 'lxml-xml'


class DOCXItemTest(unittest.TestCase):
    """DOCXText tests"""

    doc = None
    soup = None

    @classmethod
    def setUpClass(cls):
        cls.test_file_name = os.path.join('test', 'source_n1.docx')

        cls.zf = ZipFile(cls.test_file_name, 'r')

        t = cls.zf.open(DOCX_CONTENTS_FILE_NAME, 'r').read()
        cls.soup_doc = BeautifulSoup(t, DEFAULT_PARSER)

        t = cls.zf.open(DOCX_RELS_FILE_NAME, 'r').read()
        cls.soup_rels = BeautifulSoup(t, DEFAULT_PARSER)

    @classmethod
    def tearDownClass(cls):
        pass


class DOCXTextTest(DOCXItemTest):
    """DOCXText tests"""

    def test_DOCXText_getText(self):
        """getText() contents for <w:t> text element"""

        t = DOCXText(self.soup_doc.find(DOCXText.FULL_TAG_NAME))
        self.assertEqual(t.getText(), 'Бессарабов Даниил Владимирович')


class DOCXBrTest(DOCXItemTest):
    """DOCXBr tests"""

    def test_DOCXBr_getText_equals_lineseparator(self):
        """getText() for <w:br> text returns line separator"""

        t = DOCXBr(self.soup_doc.find(DOCXBr.FULL_TAG_NAME))
        self.assertEqual(t.getText(), os.linesep)


class DOCXDrawingTest(DOCXItemTest):
    """DOCXDrawing tests"""

    def setUp(self):
        # initiates DOCXDocument
        self.doc = DOCXDocument(self.test_file_name)
        self.doc.load()

        self.d_original = self.soup_doc.find(DOCXDrawing.FULL_TAG_NAME)
        self.d = DOCXDrawing(self.d_original, docx=self.doc)

    def test_DOCXDrawing_getText_returns_None(self):
        """getText() contents for <w:drawing> return None"""
        self.assertEqual(self.d.getText(), None)

    def test_DOCXDrawing_getImageName(self):
        """getImageName() contents for <w:drawing>"""
        image_name = self.d.getImageName()
        self.assertEqual(image_name, 'media/image2.jpg')


class DOCXRunTest(DOCXItemTest):
    """DOCXRun tests"""

    def setUp(self):
        self.r_original = self.soup_doc.find_all('w:br')[1].parent
        self.r = DOCXRun(self.r_original)

    def test_DOCXRun_getText_with_two_wt_children_and_one_wbr(self):
        """<w:r> with two children <w:t> and one <w:br> between them"""

        tgt = ('Родился 9 июля 1976 г. в Кемеровской области. Окончил '
               'Алтайский государственный университет (1998), Российскую '
               'Академию государственной службы при Президенте РФ (2007). '
               'Кандидат юридических наук. В 1999 г. получил статус '
               'адвоката. Работал в адвокатской конторе № 2 Железнодорожного '
               'района г. Барнаула (Алтайский край). В 2004 году избран '
               'депутатом Алтайского краевого совета народных депутатов по '
               'списку партии ЛДПР. В региональном парламенте входил в '
               'группу “Объединенные депутаты”, затем перешел во фракцию '
               '“Единая Россия”. 2 марта 2008 г. избран депутатом Алтайского '
               'краевого законодательного собрания (бывший краевой совет '
               'народных депутатов) от партии “Единая Россия”. В 2010 году '
               'назначен заместителем губернатора Алтайского края Александра '
               'Карлина (координировал деятельность органов исполнительной '
               'власти края в сфере социальной политики). С мая 2011 г. '
               'возглавлял Территориальный фонд обязательного медицинского '
               'страхования края (по должности). Президент федерации дзюдо '
               'Алтайского края. '+os.linesep+'С 2016 года - депутат '
               'Государственной Думы седьмого созыва. ')
        self.assertEqual(self.r.getText(), tgt)

    def test_DOCXRun_getCleanedText(self):
        """
        <w:r> getCleanedText() result equals
        to bs4.element.Tag.get_text()
        """
        self.assertEqual(self.r.getCleanedText(), self.r_original.get_text())


class DOCXHyperlinkTest(DOCXItemTest):
    """DOCXHyperlink tests"""

    def setUp(self):
        # initiates DOCXDocument
        self.doc = DOCXDocument(self.test_file_name)
        self.doc.load()

        # find first hyperlink
        self.h_original = self.soup_doc.find(DOCXHyperlink.FULL_TAG_NAME)
        self.h = DOCXHyperlink(self.h_original, docx=self.doc)

    def test_DOCXHyperlink_getText(self):
        """<w:hyperlink> getText"""
        tgt = ('<a href="http://www.duma.gov.ru/structure'
               '/factions/er/">ракци</a>')
        self.assertEqual(self.h.getText(), tgt)

    def test_DOCXHyperlink_getCleanedText(self):
        """
        <w:hyperlin>k getCleanedText() result
        equals to bs4.element.Tag.get_text()
        """
        self.assertEqual(self.h.getCleanedText(), self.h_original.get_text())

    def test_DOCXHyperlink_getRelationshipId(self):
        """
        <w:hyperlink> getRelationshipId returns
        correct relation identifier
        """
        self.assertEqual(self.h.getRelationshipId(), 'rId7')


class DOCXParagraphTest(DOCXItemTest):
    """DOCXParagraph tests"""

    def setUp(self):
        # initiates DOCXDocument
        self.doc = DOCXDocument(self.test_file_name)
        self.doc.load()

        # find second paragraph
        tag = DOCXParagraph.FULL_TAG_NAME
        self.p_original = self.soup_doc.find_all(tag)[1]
        self.p = DOCXParagraph(self.p_original, docx=self.doc)

        # find second paragraph
        tag = DOCXParagraph.FULL_TAG_NAME
        self.pi_original = self.soup_doc.find_all(tag)[0]
        self.pi = DOCXParagraph(self.pi_original, docx=self.doc)

    def test_DOCXParagraph_getText_with_several_children_wr(self):
        """<w:p> getText common test with several children <w:r> elements"""
        tgt = ('Депутат Государственной Думы VII созыва, избран от '
               'избирательного округа 0039 (Барнаульский - ''Алтайский край)')
        self.assertEqual(self.p.getText(), tgt)

    def test_DOCXParagraph_getId(self):
        """<w:p> getId common test"""
        self.assertEqual(self.p.getId(), '00000001')

    def test_DOCXParagraph_getChildren_doesnt_contain_excluded_tags(self):
        """<w:p> getChildren doesn't contains excluded elements"""

        tmp = [y for y in [x.name for x in self.p.getChildren()]
               if y in self.p.EXCLUDE_LIST]
        self.assertEqual(tmp, [])

    def test_DOCXParagraph_getChildren_common(self):
        """<w:p> getChildren returns first level child objects"""

        # using pi (paragraph image) which contains two child <w:r> elements
        self.assertEqual([x.name for x in self.pi.getChildren()], ['r', 'r'])

    def test_DOCXParagraph_getRawText(self):
        """<w:p> getRawText() returns list with <w:r> contents"""
        tgt = [('Депутат Государственной Думы VII созыва, избран от '
                'избирательного округа 0039 (Барнаульский '),
               '-',
               ' Алтайский край)']

        self.assertEqual(self.p.getRawText(), tgt)

    def test_DOCXParagraph_getImages_with_images_exists(self):
        """
        <w:p> getImages() return <w:drawings> only
        """
        self.assertEqual(
            set([t.name for t in self.pi.getImages()]), {'drawing'}
        )

    def test_DOCXParagraph_getImages_empty_result(self):
        """
        <w:p> getImages() return empty set if
        paragraph doesn't contain <w:drawings>
        """
        self.assertEqual([t.name for t in self.p.getImages()], [])


if __name__ == '__main__':
    unittest.main()
